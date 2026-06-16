#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Packaging_Recommendation_Overview.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x5F, 0x5F, 0x5F)
GREEN = RGBColor(0x3C, 0x6E, 0x47)
INPUT_BLUE = RGBColor(0x2F, 0x5C, 0x8A)


def set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


doc = Document()

for s in doc.sections:
    s.top_margin = Pt(40)
    s.bottom_margin = Pt(36)
    s.left_margin = Pt(50)
    s.right_margin = Pt(50)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)
style.font.color.rgb = INK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.05

# ---- Title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("AI Packaging Recommendation — Project Overview")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = INK
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sr = sub.add_run("Fine-tuned Gemma 3 (12B) model that predicts a product's full packaging "
                 "bill-of-materials from its photo and description")
sr.italic = True
sr.font.size = Pt(10)
sr.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(8)


def heading(text):
    p = doc.add_paragraph()
    rr = p.add_run(text)
    rr.bold = True
    rr.font.size = Pt(11)
    rr.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(text_runs):
    """text_runs: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(14)
    for txt, bold in text_runs:
        rn = p.add_run(txt)
        rn.bold = bold
        rn.font.size = Pt(9.5)
    return p


# ---- What we did ----
heading("What this project does")
bullet([("Synthetic dataset — ", True),
        ("we generated a synthetic dataset of ~50,000 lipstick products, each paired with its "
         "real-world packaging bill-of-materials (the boxes, tubes, caps and wraps that make it up).", False)])
bullet([("Model fine-tuning — ", True),
        ("we fine-tuned Google's Gemma 3 (12B) model on 10,000 of these data points, teaching it to "
         "predict a complete, structured packaging breakdown for a given product.", False)])
bullet([("Benchmarking — ", True),
        ("we manually selected 30 real lipstick products from mecca.com and ran them through the "
         "model to generate packaging recommendations. The attached spreadsheet holds those 30 results.", False)])

# ---- How to read the sheet ----
heading("How to read the spreadsheet")
intro = doc.add_paragraph()
ir = intro.add_run("The sheet has two parts. The ")
ir.font.size = Pt(9.5)
b1 = intro.add_run("INPUT")
b1.bold = True
b1.font.color.rgb = INPUT_BLUE
b1.font.size = Pt(9.5)
m1 = intro.add_run(" columns are what we gave the model; the ")
m1.font.size = Pt(9.5)
b2 = intro.add_run("LARGE LANGUAGE MODEL OUTPUT")
b2.bold = True
b2.font.color.rgb = GREEN
b2.font.size = Pt(9.5)
m2 = intro.add_run(" columns are what the model produced. Each product spans several rows — "
                   "one row per packaging component.")
m2.font.size = Pt(9.5)
intro.paragraph_format.space_after = Pt(5)

# ---- Column reference table ----
ROWS = [
    ("INPUT", "Product", "The product / brand name."),
    ("INPUT", "Link", "The mecca.com product page the item was taken from."),
    ("INPUT", "Image caption", "A plain description of the product photo, produced by the model from the image."),
    ("OUTPUT", "Packaging line", "Which packaging tier the part belongs to: PP = primary (touches the product), "
                                 "SP = secondary (retail box), TP = tertiary (shipping/transport)."),
    ("OUTPUT", "Component", "The individual packaging part (e.g. lipstick tube, cap, carton box, pallet)."),
    ("OUTPUT", "Rigid / soft", "Whether the part is rigid (hard) or soft (flexible)."),
    ("OUTPUT", "Is reusable", "Whether the part is designed to be refilled or reused."),
    ("OUTPUT", "Length / Width / Height (mm)", "Estimated outer dimensions of the part, in millimetres."),
    ("OUTPUT", "Shape", "The part's inferred form: box, cylindrical, or flat."),
    ("OUTPUT", "Volume (cm³)", "Estimated volume of the part, derived from its dimensions."),
    ("OUTPUT", "Materials", "The material(s) the part is made of, each with its mass (e.g. Polypropylene:8.2g)."),
    ("OUTPUT", "Mass (g)", "Total weight of the part, in grams."),
    ("OUTPUT", "Carbon footprint (kg)", "Estimated CO₂ emissions for the part, looked up from a materials catalogue."),
    ("OUTPUT", "Water consumption", "Estimated water use for the part, looked up from the same catalogue."),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Light Grid Accent 3"
table.autofit = True

hdr = table.rows[0].cells
for i, label in enumerate(("Part", "Column", "What it means")):
    hdr[i].text = ""
    pr = hdr[i].paragraphs[0]
    rn = pr.add_run(label)
    rn.bold = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(hdr[i], "3C6E47")

for part, col, desc in ROWS:
    cells = table.add_row().cells
    # Part
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(part)
    r0.bold = True
    r0.font.size = Pt(8.5)
    r0.font.color.rgb = INPUT_BLUE if part == "INPUT" else GREEN
    set_cell_bg(cells[0], "EAF1FB" if part == "INPUT" else "EAF3EC")
    # Column
    p1 = cells[1].paragraphs[0]
    r1 = p1.add_run(col)
    r1.bold = True
    r1.font.size = Pt(8.5)
    # Description
    p2 = cells[2].paragraphs[0]
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8.5)
    for c in cells:
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].paragraph_format.space_before = Pt(0)

# ---- footer note ----
note = doc.add_paragraph()
nr = note.add_run("Note: carbon and water figures are not guessed by the model — they are calculated "
                  "by matching each predicted material to a real environmental catalogue, then scaling by mass.")
nr.italic = True
nr.font.size = Pt(8)
nr.font.color.rgb = GREY
note.paragraph_format.space_before = Pt(6)

doc.save(OUT)
print(f"wrote {OUT}")
